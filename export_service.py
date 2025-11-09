"""
Export Service - PDF, EPUB, DOCX generation
"""
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER
from ebooklib import epub
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from datetime import datetime


class ExportService:
    """Service for exporting translations"""
    
    def __init__(self):
        self.output_dir = "exports"
        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir)
    
    def export_to_pdf(self, project_name: str, chapters: list, metadata: dict = None) -> str:
        """Export chapters to PDF"""
        filename = f"{project_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        filepath = os.path.join(self.output_dir, filename)
        
        # Create PDF
        doc = SimpleDocTemplate(
            filepath,
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18
        )
        
        # Container for 'Flowable' objects
        story = []
        
        # Define styles
        styles = getSampleStyleSheet()
        
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor='#4f46e5',
            spaceAfter=30,
            alignment=TA_CENTER
        )
        
        chapter_title_style = ParagraphStyle(
            'ChapterTitle',
            parent=styles['Heading2'],
            fontSize=16,
            textColor='#6366f1',
            spaceAfter=12,
            spaceBefore=12
        )
        
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['BodyText'],
            fontSize=11,
            alignment=TA_JUSTIFY,
            spaceAfter=12,
            leading=16
        )
        
        # Add title
        story.append(Paragraph(project_name, title_style))
        story.append(Spacer(1, 12))
        
        # Add metadata if provided
        if metadata:
            info_style = styles['Normal']
            if metadata.get('description'):
                story.append(Paragraph(f"<i>{metadata['description']}</i>", info_style))
            story.append(Paragraph(f"<b>Kaynak Dil:</b> {metadata.get('source_language', 'N/A')}", info_style))
            story.append(Paragraph(f"<b>Hedef Dil:</b> {metadata.get('target_language', 'N/A')}", info_style))
            story.append(Paragraph(f"<b>Bölüm Sayısı:</b> {len(chapters)}", info_style))
            story.append(Paragraph(f"<b>Oluşturulma:</b> {datetime.now().strftime('%d.%m.%Y %H:%M')}", info_style))
        
        story.append(Spacer(1, 24))
        story.append(PageBreak())
        
        # Add chapters
        for chapter in chapters:
            chapter_num = chapter.get('chapter_number', 'N/A')
            title = chapter.get('title', '')
            text = chapter.get('translated_text', '')
            
            # Chapter title
            chapter_heading = f"Bölüm {chapter_num}"
            if title:
                chapter_heading += f": {title}"
            
            story.append(Paragraph(chapter_heading, chapter_title_style))
            story.append(Spacer(1, 12))
            
            # Chapter text
            paragraphs = text.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    story.append(Paragraph(para.strip(), body_style))
                    story.append(Spacer(1, 6))
            
            story.append(PageBreak())
        
        # Build PDF
        doc.build(story)
        
        return filepath
    
    def export_to_epub(self, project_name: str, chapters: list, metadata: dict = None) -> str:
        """Export chapters to EPUB"""
        filename = f"{project_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.epub"
        filepath = os.path.join(self.output_dir, filename)
        
        # Create EPUB book
        book = epub.EpubBook()
        
        # Set metadata
        book.set_identifier(f'novel_translator_{datetime.now().timestamp()}')
        book.set_title(project_name)
        book.set_language(metadata.get('target_language', 'tr'))
        book.add_author('Novel Translator')
        
        # Add chapters
        epub_chapters = []
        spine = ['nav']
        
        for idx, chapter in enumerate(chapters):
            chapter_num = chapter.get('chapter_number', idx + 1)
            title = chapter.get('title', '')
            text = chapter.get('translated_text', '')
            
            # Create chapter
            chapter_title = f"Bölüm {chapter_num}"
            if title:
                chapter_title += f": {title}"
            
            epub_chapter = epub.EpubHtml(
                title=chapter_title,
                file_name=f'chapter_{chapter_num}.xhtml',
                lang=metadata.get('target_language', 'tr')
            )
            
            # Format content
            content = f'<h1>{chapter_title}</h1>'
            paragraphs = text.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    content += f'<p>{para.strip()}</p>'
            
            epub_chapter.content = content
            
            book.add_item(epub_chapter)
            epub_chapters.append(epub_chapter)
            spine.append(epub_chapter)
        
        # Add default NCX and Nav
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())
        
        # Define CSS
        style = '''
        @namespace epub "http://www.idpf.org/2007/ops";
        body {
            font-family: Georgia, serif;
            line-height: 1.8;
        }
        h1 {
            color: #4f46e5;
            text-align: center;
            margin-bottom: 2em;
        }
        p {
            text-align: justify;
            margin: 1em 0;
        }
        '''
        nav_css = epub.EpubItem(
            uid="style_nav",
            file_name="style/nav.css",
            media_type="text/css",
            content=style
        )
        book.add_item(nav_css)
        
        # Define Table of Contents
        book.toc = epub_chapters
        
        # Add spine
        book.spine = spine
        
        # Write EPUB file
        epub.write_epub(filepath, book)
        
        return filepath
    
    def export_to_docx(self, project_name: str, chapters: list, metadata: dict = None) -> str:
        """Export chapters to DOCX"""
        filename = f"{project_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = os.path.join(self.output_dir, filename)
        
        # Create document
        doc = Document()
        
        # Set document properties
        core_properties = doc.core_properties
        core_properties.title = project_name
        core_properties.author = 'Novel Translator'
        core_properties.created = datetime.now()
        
        # Add title
        title = doc.add_heading(project_name, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata
        if metadata:
            if metadata.get('description'):
                p = doc.add_paragraph(metadata['description'])
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.runs[0].italic = True
            
            doc.add_paragraph()
            doc.add_paragraph(f"Kaynak Dil: {metadata.get('source_language', 'N/A')}")
            doc.add_paragraph(f"Hedef Dil: {metadata.get('target_language', 'N/A')}")
            doc.add_paragraph(f"Bölüm Sayısı: {len(chapters)}")
            doc.add_paragraph(f"Oluşturulma: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
        
        doc.add_page_break()
        
        # Add chapters
        for chapter in chapters:
            chapter_num = chapter.get('chapter_number', 'N/A')
            title = chapter.get('title', '')
            text = chapter.get('translated_text', '')
            
            # Chapter title
            chapter_heading = f"Bölüm {chapter_num}"
            if title:
                chapter_heading += f": {title}"
            
            doc.add_heading(chapter_heading, 1)
            
            # Chapter text
            paragraphs = text.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    p = doc.add_paragraph(para.strip())
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    
                    # Set font
                    for run in p.runs:
                        run.font.size = Pt(11)
                        run.font.name = 'Georgia'
            
            doc.add_page_break()
        
        # Save document
        doc.save(filepath)
        
        return filepath
    
    def export_to_txt(self, project_name: str, chapters: list, metadata: dict = None) -> str:
        """Export chapters to plain text"""
        filename = f"{project_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
        filepath = os.path.join(self.output_dir, filename)
        
        with open(filepath, 'w', encoding='utf-8') as f:
            # Title
            f.write(f"{'=' * 60}\n")
            f.write(f"{project_name.center(60)}\n")
            f.write(f"{'=' * 60}\n\n")
            
            # Metadata
            if metadata:
                if metadata.get('description'):
                    f.write(f"{metadata['description']}\n\n")
                f.write(f"Kaynak Dil: {metadata.get('source_language', 'N/A')}\n")
                f.write(f"Hedef Dil: {metadata.get('target_language', 'N/A')}\n")
                f.write(f"Bölüm Sayısı: {len(chapters)}\n")
                f.write(f"Oluşturulma: {datetime.now().strftime('%d.%m.%Y %H:%M')}\n\n")
            
            f.write(f"{'=' * 60}\n\n")
            
            # Chapters
            for chapter in chapters:
                chapter_num = chapter.get('chapter_number', 'N/A')
                title = chapter.get('title', '')
                text = chapter.get('translated_text', '')
                
                # Chapter title
                chapter_heading = f"Bölüm {chapter_num}"
                if title:
                    chapter_heading += f": {title}"
                
                f.write(f"\n\n{chapter_heading}\n")
                f.write(f"{'-' * len(chapter_heading)}\n\n")
                
                # Chapter text
                f.write(text)
                f.write("\n\n")
        
        return filepath

